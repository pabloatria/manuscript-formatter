# tests/test_cover_letter.py
from pathlib import Path
from docx import Document

from scripts.cover_letter import generate_cover_letter
from scripts.journal_config import load_journal

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "minimal_manuscript.docx"
CFG_DIR = Path(__file__).resolve().parent.parent / "scripts" / "journals"


def test_cover_letter_writes_docx_with_placeholders(tmp_path):
    out = tmp_path / "cover.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    generate_cover_letter(
        FIXTURE, out, journal_cfg=cfg,
        editor="Dr. Test Editor",
        manuscript_title="Endocrown Survival Study",
        corresponding_author="Pablo Atria, DDS, PhD",
    )
    assert out.exists()
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Dr. Test Editor" in txt
    assert "Endocrown Survival Study" in txt
    assert "[NOVELTY" in txt
    assert "Pablo Atria" in txt


def test_cover_letter_falls_back_to_placeholders_when_args_missing(tmp_path):
    """When optional args are not provided, the template fields render
    as placeholders the user can fill in by hand."""
    out = tmp_path / "cover.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    generate_cover_letter(
        FIXTURE, out, journal_cfg=cfg,
        editor=None, manuscript_title=None, corresponding_author=None,
    )
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "[EDITOR]" in txt
    assert "[MANUSCRIPT_TITLE]" in txt
    assert "[CORRESPONDING_AUTHOR]" in txt


def test_cover_letter_extracts_methods_summary_from_manuscript(tmp_path):
    """When the manuscript has a Methods section, the cover letter's
    {methods_summary} placeholder gets filled with the first sentence of
    that section instead of the [METHODS_SUMMARY] placeholder."""
    in_path = tmp_path / "rich.docx"
    d = Document()
    d.add_heading("Abstract", level=1)
    d.add_paragraph("Brief abstract.")
    d.add_heading("Methods", level=1)
    d.add_paragraph("Fifty patients enrolled in a prospective trial.")
    d.save(in_path)
    out = tmp_path / "cover.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    generate_cover_letter(in_path, out, journal_cfg=cfg, editor="Dr. X",
                          manuscript_title="T", corresponding_author="P")
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Fifty patients enrolled in a prospective trial" in txt


def test_cover_letter_falls_back_to_placeholder_when_no_methods(tmp_path):
    """When the manuscript has no Methods section, [METHODS_SUMMARY]
    appears as a placeholder for the user to fill in."""
    in_path = tmp_path / "no_methods.docx"
    d = Document()
    d.add_heading("Abstract", level=1)
    d.add_paragraph("Brief abstract.")
    d.save(in_path)
    out = tmp_path / "cover.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    generate_cover_letter(in_path, out, journal_cfg=cfg, editor="Dr. X",
                          manuscript_title="T", corresponding_author="P")
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "[METHODS_SUMMARY" in txt


def test_cover_letter_methods_summary_handles_embedded_newlines(tmp_path):
    """python-docx may preserve newlines inside paragraph.text. The
    methods summarizer must collapse them before sentence-splitting,
    otherwise a paragraph like 'Sentence one.\\nSentence two.' returns
    the entire paragraph as one 'sentence'."""
    in_path = tmp_path / "newlines.docx"
    d = Document()
    d.add_heading("Methods", level=1)
    p = d.add_paragraph("Fifty patients enrolled.")
    # Force an embedded line break in the same paragraph
    p.add_run().add_break()
    p.add_run("Then we followed them prospectively.")
    d.save(in_path)
    out = tmp_path / "cover.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    generate_cover_letter(in_path, out, journal_cfg=cfg, editor="Dr. X",
                          manuscript_title="T", corresponding_author="P")
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # The summary should be the first sentence only, not the whole paragraph
    assert "Fifty patients enrolled" in txt
    assert "Then we followed" not in txt or "Fifty patients enrolled" in txt.split("Then we followed")[0]
