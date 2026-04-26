from pathlib import Path

import pytest
from docx import Document

from scripts.docx_helpers import read_headings, Heading

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "minimal_manuscript.docx"


def test_read_headings_finds_six_top_level():
    hs = read_headings(FIXTURE)
    assert all(isinstance(h, Heading) for h in hs)
    assert all(h.level == 1 for h in hs)
    assert [h.text for h in hs] == ["Abstract", "Background", "Methods", "Results",
                                    "Discussion", "Conclusion"]


def test_each_heading_has_paragraph_index():
    hs = read_headings(FIXTURE)
    indices = [h.paragraph_index for h in hs]
    assert indices == sorted(indices)
    assert indices[0] >= 0


def test_read_headings_empty_document(tmp_path):
    p = tmp_path / "empty.docx"
    Document().save(p)
    assert read_headings(p) == []


def test_read_headings_no_headings(tmp_path):
    p = tmp_path / "body_only.docx"
    d = Document()
    d.add_paragraph("Just body text.")
    d.add_paragraph("More body text.")
    d.save(p)
    assert read_headings(p) == []


def test_read_headings_preserves_nested_levels(tmp_path):
    """Catches the silent-regression risk where _heading_level always returns 1."""
    p = tmp_path / "nested.docx"
    d = Document()
    d.add_heading("Methods", level=1)
    d.add_heading("Participants", level=2)
    d.add_heading("Inclusion criteria", level=3)
    d.save(p)
    hs = read_headings(p)
    assert [h.level for h in hs] == [1, 2, 3]
