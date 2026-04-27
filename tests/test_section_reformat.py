# tests/test_section_reformat.py
from pathlib import Path
from docx import Document
import pytest

from scripts.docx_helpers import reformat_sections, read_headings
from scripts.journal_config import load_journal

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "minimal_manuscript.docx"
CFG_DIR = Path(__file__).resolve().parent.parent / "scripts" / "journals"


def _body_paragraph_texts(path: Path) -> list[str]:
    """All non-heading paragraph texts (the author's prose)."""
    doc = Document(str(path))
    out = []
    for p in doc.paragraphs:
        style = p.style.name or ""
        if not style.startswith("Heading "):
            out.append(p.text)
    return out


def test_reformat_renames_headings_to_jpd_canonical(tmp_path):
    out_path = tmp_path / "out.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    reformat_sections(FIXTURE, out_path, cfg)
    new_headings = [h.text for h in read_headings(out_path)]
    # JPD canonicals: Abstract, Introduction, Material and Methods,
    # Results, Discussion, Conclusions
    assert "Introduction" in new_headings          # was "Background"
    assert "Material and Methods" in new_headings  # was "Methods"
    assert "Conclusions" in new_headings           # was "Conclusion"
    # Headings that didn't need renaming should still be present
    assert "Abstract" in new_headings
    assert "Results" in new_headings
    assert "Discussion" in new_headings


def test_reformat_preserves_body_paragraph_text(tmp_path):
    """Body paragraphs' visible text must be byte-identical between input
    and output. This is a text-equality check (not full XML byte-equality —
    the latter is Task 18's parametrized invariant)."""
    out_path = tmp_path / "out.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    reformat_sections(FIXTURE, out_path, cfg)
    assert _body_paragraph_texts(FIXTURE) == _body_paragraph_texts(out_path), \
        "BODY-TEXT PRESERVATION VIOLATED — body prose changed"


def test_reformat_preserves_paragraph_count(tmp_path):
    """Every paragraph in the input must appear in the output, in order."""
    out_path = tmp_path / "out.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    reformat_sections(FIXTURE, out_path, cfg)
    in_doc = Document(str(FIXTURE))
    out_doc = Document(str(out_path))
    assert len(in_doc.paragraphs) == len(out_doc.paragraphs)


def test_unmapped_heading_is_left_alone(tmp_path):
    """A heading the JPD config doesn't recognize must keep its original text.
    We construct a manuscript with an 'Acknowledgments' heading that no JPD
    section claims; reformat must not rename it."""
    in_path = tmp_path / "in.docx"
    d = Document()
    d.add_heading("Abstract", level=1)
    d.add_paragraph("Brief abstract.")
    d.add_heading("Acknowledgments", level=1)  # not in JPD aliases
    d.add_paragraph("Thanks to the team.")
    d.save(in_path)
    out_path = tmp_path / "out.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    reformat_sections(in_path, out_path, cfg)
    out_headings = [h.text for h in read_headings(out_path)]
    assert "Acknowledgments" in out_headings
    assert "Abstract" in out_headings


def test_in_place_rewrite_is_rejected(tmp_path):
    """Calling with input_path == output_path must raise rather than risk
    a non-atomic rewrite that could corrupt the input."""
    p = tmp_path / "same.docx"
    Document().save(p)
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    with pytest.raises(ValueError, match="in-place rewrite"):
        reformat_sections(p, p, cfg)


def test_orphan_empty_runs_are_dropped_after_rename(tmp_path):
    """A heading with two runs (e.g., bold prefix + plain suffix) should
    have only one run after rename, not one with the new text and a stale
    empty second run."""
    in_path = tmp_path / "multirun.docx"
    d = Document()
    h = d.add_heading("", level=1)
    # Two runs: simulate "Back" (bold) + "ground" (plain)
    r1 = h.add_run("Back")
    r1.bold = True
    h.add_run("ground")
    d.save(in_path)
    out_path = tmp_path / "out.docx"
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    reformat_sections(in_path, out_path, cfg)
    out_doc = Document(out_path)
    # Find the renamed heading in output
    renamed = next(p for p in out_doc.paragraphs
                   if (p.style.name or "").startswith("Heading "))
    assert len(renamed.runs) == 1
    assert renamed.runs[0].text == "Introduction"
