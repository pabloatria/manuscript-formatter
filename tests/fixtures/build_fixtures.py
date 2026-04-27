"""Generate a small .docx fixture for tests. Run once, commit the .docx.

After python-docx writes the file, we rewrite the ZIP container so that
ZIP modtimes are fixed at 2000-01-01 and the dcterms:created/modified
timestamps in docProps/core.xml are pinned to the same epoch. This makes
the binary byte-stable across rebuilds — important for Task 18's
text-preservation invariant test.
"""
import re
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

EPOCH_ISO = (datetime(2000, 1, 1, tzinfo=timezone.utc)
             .isoformat().replace("+00:00", "Z"))
EPOCH_ZIP = (2000, 1, 1, 0, 0, 0)


def _normalize_docx(path: Path) -> None:
    """Zero out volatile metadata so the fixture is byte-stable."""
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td) / "out.docx"
        with zipfile.ZipFile(path) as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in sorted(zin.infolist(), key=lambda i: i.filename):
                data = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    data = re.sub(
                        rb"<dcterms:created[^>]*>[^<]*</dcterms:created>",
                        f'<dcterms:created xsi:type="dcterms:W3CDTF">'
                        f"{EPOCH_ISO}</dcterms:created>".encode(),
                        data,
                    )
                    data = re.sub(
                        rb"<dcterms:modified[^>]*>[^<]*</dcterms:modified>",
                        f'<dcterms:modified xsi:type="dcterms:W3CDTF">'
                        f"{EPOCH_ISO}</dcterms:modified>".encode(),
                        data,
                    )
                new_info = zipfile.ZipInfo(info.filename, date_time=EPOCH_ZIP)
                new_info.compress_type = zipfile.ZIP_DEFLATED
                zout.writestr(new_info, data)
        shutil.move(str(tmp), str(path))


def make_jpd_style_manuscript(out_path: Path):
    doc = Document()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Statement of Problem: Endocrowns are minimally invasive but their "
        "long-term survival in posterior teeth remains underreported. "
        "Purpose: to evaluate 5-year survival of lithium disilicate endocrowns. "
        "Methods: 50 endocrowns followed prospectively. Results: 92% survival. "
        "Conclusions: comparable to crowns."
    )
    doc.add_heading("Background", level=1)  # alias for "Statement of Problem"
    doc.add_paragraph("Endodontically treated teeth are restorative challenges...")
    doc.add_heading("Methods", level=1)  # alias for "Material and Methods"
    doc.add_paragraph("Fifty patients enrolled, ages 35-72...")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Kaplan-Meier 5-year survival 92%...")
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("Our findings agree with prior systematic reviews...")
    doc.add_heading("Conclusion", level=1)  # alias for "Conclusions"
    doc.add_paragraph("Endocrowns are a reliable alternative.")
    doc.save(out_path)
    _normalize_docx(out_path)


def make_jpd_clinical_report(out_path: Path):
    """A short clinical report following JPD's required structure:
    Abstract → Introduction → Clinical Report → Discussion → Summary."""
    doc = Document()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "A 62-year-old patient with a failing implant-supported prosthesis "
        "was treated with a digital workflow combining intraoral scanning, "
        "guided implant placement, and a milled zirconia crown. Six-month "
        "follow-up showed stable peri-implant health and patient satisfaction."
    )
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Failures of implant-supported prostheses in the maxillary anterior "
        "are challenging because of esthetic demands and bone quality."
    )
    doc.add_heading("Clinical Report", level=1)
    doc.add_paragraph(
        "Patient history: A 62-year-old female presented with an exposed "
        "abutment screw on tooth 11. Examination revealed peri-implantitis "
        "with 4 mm probing depth and bleeding on probing."
    )
    doc.add_paragraph(
        "Treatment: After non-surgical debridement and a 6-week healing "
        "period, the implant was explanted using a guided technique."
    )
    doc.add_paragraph(
        "Outcome: At 6-month follow-up, the new implant showed full "
        "osseointegration with stable peri-implant health."
    )
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "Digital workflows for implant explantation and replacement reduce "
        "surgical time and improve esthetic predictability."
    )
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "This clinical report demonstrates a digital workflow for the "
        "predictable replacement of a failing maxillary anterior implant."
    )
    doc.save(out_path)
    _normalize_docx(out_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "minimal_manuscript.docx"
    make_jpd_style_manuscript(out)
    print(f"wrote {out}")

    out = Path(__file__).resolve().parent / "clinical_report_manuscript.docx"
    make_jpd_clinical_report(out)
    print(f"wrote {out}")
