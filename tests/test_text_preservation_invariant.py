import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from xml.etree import ElementTree as ET

REPO = Path(__file__).resolve().parent.parent
FIXT = REPO / "tests" / "fixtures"
CLI = REPO / "scripts" / "format_manuscript.py"


# (slug, article_type, fixture_filename) — only populated combinations
PARAMETER_GRID: list[tuple[str, str, str]] = []
for slug in ["jpd", "jerd", "jomi", "coir", "int-j-prosthodont"]:
    for at in ["research", "case-report", "technique", "systematic-review"]:
        fixture = ("clinical_report_manuscript.docx" if at == "case-report"
                   else "minimal_manuscript.docx")
        PARAMETER_GRID.append((slug, at, fixture))
# research-only journals
for slug in ["jdr", "jada", "j-dent", "j-periodontol", "j-endod", "oper-dent"]:
    PARAMETER_GRID.append((slug, "research", "minimal_manuscript.docx"))


def _run_cli(in_path: Path, slug: str, article_type: str, out_dir: Path) -> Path:
    cmd = [
        sys.executable, str(CLI), str(in_path),
        "--references", str(FIXT / "sample_zotero.json"),
        "--journal", slug,
        "--article-type", article_type,
        "--out-dir", str(out_dir),
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    assert r.returncode == 0, (
        f"CLI failed for {slug}/{article_type}:\n{r.stderr}"
    )
    return out_dir / f"{in_path.stem}-{slug}.docx"


def _body_paragraph_texts(p: Path) -> list[str]:
    doc = Document(str(p))
    return [par.text for par in doc.paragraphs
            if not (par.style.name or "").startswith("Heading ")]


def _body_paragraph_xml(p: Path) -> list[bytes]:
    doc = Document(str(p))
    return [ET.tostring(par._element, encoding="utf-8")
            for par in doc.paragraphs
            if not (par.style.name or "").startswith("Heading ")]


@pytest.mark.parametrize("slug,article_type,fixture", PARAMETER_GRID)
def test_body_text_byte_identical(slug, article_type, fixture, tmp_path):
    in_path = FIXT / fixture
    out_path = _run_cli(in_path, slug, article_type, tmp_path)
    assert _body_paragraph_texts(in_path) == _body_paragraph_texts(out_path), \
        f"{slug}/{article_type}: BODY-TEXT INVARIANT VIOLATED"


@pytest.mark.parametrize("slug,article_type,fixture", PARAMETER_GRID)
def test_body_paragraph_xml_byte_identical(slug, article_type, fixture, tmp_path):
    in_path = FIXT / fixture
    out_path = _run_cli(in_path, slug, article_type, tmp_path)
    assert _body_paragraph_xml(in_path) == _body_paragraph_xml(out_path), \
        f"{slug}/{article_type}: BODY-XML INVARIANT VIOLATED"
