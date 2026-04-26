# tests/test_validators.py
from pathlib import Path
from docx import Document

from scripts.validators import validate_manuscript
from scripts.journal_config import load_journal

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "minimal_manuscript.docx"
CFG_DIR = Path(__file__).resolve().parent.parent / "scripts" / "journals"


def test_validator_returns_section_word_counts():
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(FIXTURE, cfg)
    by_canon = {r["canonical"]: r for r in report["sections"] if r["canonical"]}
    assert by_canon["abstract"]["word_count"] > 0
    assert by_canon["abstract"]["word_count"] < 250  # fixture is short


def test_validator_flags_missing_required_section():
    """JPD requires a Purpose section. The fixture has no Purpose, so it
    must be flagged as missing."""
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(FIXTURE, cfg)
    missing = [m["canonical"] for m in report["missing_required"]]
    assert "purpose" in missing


def test_validator_flags_overlong_abstract(tmp_path):
    """Synthesize a manuscript whose abstract exceeds 250 words; expect
    over_limit=True with limit=250."""
    in_path = tmp_path / "long_abstract.docx"
    doc = Document()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("word " * 300)
    doc.save(in_path)
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(in_path, cfg)
    abs_row = next(r for r in report["sections"] if r["canonical"] == "abstract")
    assert abs_row["over_limit"] is True
    assert abs_row["limit"] == 250


def test_validator_returns_total_word_count():
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(FIXTURE, cfg)
    assert "total_word_count" in report
    assert report["total_word_count"] > 0
    assert report["total_word_limit"] == 5000
    assert report["total_over_limit"] is False  # fixture is small


def test_validator_includes_original_heading_for_audit():
    """The fixture has 'Background' which should map to canonical
    'introduction'. The report should preserve the original heading text
    so the audit trail in Task 14 can show 'Background renamed -> Statement of Problem'."""
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(FIXTURE, cfg)
    intro_row = next(r for r in report["sections"] if r["canonical"] == "introduction")
    assert intro_row["original_heading"] == "Background"


def test_total_word_count_excludes_abstract(tmp_path):
    """Build a manuscript with a 200-word abstract and 300-word methods.
    Total should be 300, not 500 (abstract is excluded from main-text total)."""
    in_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("word " * 200)
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("methods-word " * 300)
    doc.save(in_path)
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(in_path, cfg)
    assert report["total_word_count"] == 300


def test_missing_required_returns_canonical_and_display_only(tmp_path):
    """missing_required entries must be {canonical, display}, no extras."""
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(FIXTURE, cfg)
    for entry in report["missing_required"]:
        assert set(entry.keys()) == {"canonical", "display"}


def test_validator_captures_prose_before_first_heading(tmp_path):
    """Prose before the first heading must show up as a pseudo-block in
    sections, not be silently lost."""
    in_path = tmp_path / "lead.docx"
    doc = Document()
    doc.add_paragraph("Some unstyled prose before any heading.")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("Brief abstract.")
    doc.save(in_path)
    cfg = load_journal("jpd", config_dir=CFG_DIR)
    report = validate_manuscript(in_path, cfg)
    pseudo = next(
        (r for r in report["sections"] if r["original_heading"] == "(before first heading)"),
        None,
    )
    assert pseudo is not None
    assert pseudo["canonical"] is None
    assert pseudo["word_count"] > 0
