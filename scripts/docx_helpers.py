"""python-docx helpers: heading detection, paragraph manipulation."""
from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.text.paragraph import Paragraph
from rapidfuzz import fuzz

# Match exactly "Heading 1" through "Heading 9" — refuses "Heading 1 Char"
# (Word's auto-generated linked character style), "Heading 12", localized
# names like "Encabezado 1", and any non-Latin variants. We deliberately
# do NOT silently downgrade or guess at level — unrecognized headings
# return None and the caller sees them as body paragraphs.
_HEADING_RE = re.compile(r"^Heading ([1-9])$")


@dataclass(frozen=True)
class Heading:
    level: int
    text: str
    paragraph_index: int


def _heading_level(p: Paragraph) -> int | None:
    """Return 1..9 if the paragraph uses a Heading 1..9 style; else None.

    Tightens the previous startswith() check: rejects "Heading 1 Char" and
    "Heading 12" rather than relying on a bare-except to silently drop
    them.
    """
    style_name = getattr(p.style, "name", "") or ""
    m = _HEADING_RE.match(style_name)
    return int(m.group(1)) if m else None


def read_headings(docx_path: Path) -> list[Heading]:
    """Return the document's heading paragraphs in order.

    Each entry is a Heading dataclass with level (1..9), stripped text, and
    the absolute paragraph_index in `Document(docx_path).paragraphs`.
    """
    doc = Document(docx_path)
    out = []
    for idx, p in enumerate(doc.paragraphs):
        lv = _heading_level(p)
        if lv is not None:
            out.append(Heading(level=lv, text=p.text.strip(), paragraph_index=idx))
    return out


# Below this score (0..100), a heading is treated as unmapped (canonical=None).
# 82 chosen to exclude false-positives like fuzz.ratio("Funding", "Findings")
# (which scores exactly 80) while keeping legitimate fuzzy matches above
# threshold (e.g., "Conclussion" -> "Conclusions" scores 90.9).
FUZZY_THRESHOLD = 82


@dataclass(frozen=True)
class MappedHeading:
    """A detected heading paired with its journal-canonical mapping.

    `text` is the heading content as it currently exists. `original_text` is
    the verbatim string read from the input document and is preserved across
    any future normalization (case, whitespace, punctuation) that downstream
    code might apply to `text`. Today the two are identical; the split exists
    so the audit trail in the validation report can always show what the
    user actually wrote, even after section-rename has overwritten `text`
    with the canonical form.
    """
    level: int
    text: str
    paragraph_index: int
    canonical: str | None
    confidence: float
    original_text: str


def map_headings_to_canonical(
    headings: list[Heading],
    sections_config: list[dict],
) -> list[MappedHeading]:
    """For each detected heading, find the best canonical match across all
    aliases declared in the journal config.

    Returns parallel MappedHeading list. Unmapped headings keep
    canonical=None and original_text preserved so the caller can leave them
    untouched in the output document.
    """
    out: list[MappedHeading] = []
    for h in headings:
        best_canon: str | None = None
        best_score = 0.0
        for sec in sections_config:
            for alias in sec.get("aliases", []):
                score = fuzz.ratio(h.text.lower(), alias.lower())
                if score > best_score:
                    best_score = score
                    best_canon = sec["canonical"]
        out.append(MappedHeading(
            level=h.level,
            text=h.text,
            paragraph_index=h.paragraph_index,
            canonical=best_canon if best_score >= FUZZY_THRESHOLD else None,
            confidence=best_score / 100.0,
            original_text=h.text,
        ))
    return out
