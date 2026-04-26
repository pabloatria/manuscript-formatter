"""Generate a cover-letter .docx from a journal-config template.

The template is journal_cfg["cover_letter_template"] — a multiline string
with {placeholder} fields. We fill in what we can (editor, title,
corresponding author, methods summary extracted from the manuscript) and
leave clearly-labeled placeholders for the rest (notably [NOVELTY] which
requires the author's editorial judgment about why this manuscript is a
fit for this journal).
"""
from pathlib import Path

from docx import Document

from .docx_helpers import read_headings


def _summarize_methods(input_docx: Path) -> str:
    """Return one sentence describing the manuscript's methods.

    Looks for a heading whose stripped text is one of the common Methods
    aliases ("Methods", "Material and Methods", "Materials and Methods")
    and returns the first sentence of the first non-empty body paragraph
    underneath it. If no Methods section is found, returns the
    [METHODS_SUMMARY] placeholder so the user knows to fill it in.
    """
    aliases = {
        "methods", "material and methods", "materials and methods",
    }
    headings = read_headings(input_docx)
    target_idx = None
    for h in headings:
        if h.text.strip().lower() in aliases:
            target_idx = h.paragraph_index
            break
    if target_idx is None:
        return "[METHODS_SUMMARY: one-sentence description of the study design]"

    doc = Document(input_docx)
    paras = doc.paragraphs
    for j in range(target_idx + 1, len(paras)):
        style = paras[j].style.name or ""
        if style.startswith("Heading "):
            break
        text = " ".join(paras[j].text.split())  # collapse all whitespace
        if text:
            # Naive sentence split on '. ' — abbreviations like "Dr." inside
            # a sentence may produce truncated output. The cover letter is a
            # draft the user reviews before sending, so this is acceptable.
            first = text.split(". ")[0].rstrip(".")
            return first + "."
    return "[METHODS_SUMMARY: one-sentence description of the study design]"


def generate_cover_letter(
    input_docx: Path,
    output_docx: Path,
    journal_cfg: dict,
    editor: str | None = None,
    manuscript_title: str | None = None,
    corresponding_author: str | None = None,
) -> None:
    """Render the journal's cover_letter_template with available fields
    substituted; leave bracketed placeholders where the user must fill in.
    """
    template = journal_cfg.get("cover_letter_template", "")
    body = template.format(
        editor=editor or "[EDITOR]",
        title=manuscript_title or "[MANUSCRIPT_TITLE]",
        methods_summary=_summarize_methods(input_docx),
        results_summary="[RESULTS_SUMMARY: one sentence on the headline finding]",
        authors_summary="The authors have all contributed substantially to this work",
        corresponding_author=corresponding_author or "[CORRESPONDING_AUTHOR]",
    )
    doc = Document()
    for paragraph_text in body.strip().split("\n\n"):
        doc.add_paragraph(paragraph_text.strip())
    doc.save(output_docx)
